"""
PdfSuite Web Application - FastAPI Backend
All PDF operations are performed entirely in-memory using BytesIO streams.
No files are saved to disk.
"""

import io
import zipfile
from typing import List

from fastapi import FastAPI, File, UploadFile, HTTPException, Request, Form
from fastapi.responses import StreamingResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates

# PDF Libraries
from pypdf import PdfReader, PdfWriter
import fitz  # PyMuPDF - for compress, pdf->jpg
from PIL import Image
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas
from reportlab.platypus import Table, TableStyle

# --- App Setup ---
app = FastAPI(title="PdfSuite", description="PDF Utilities Web App")

# Templates for HTML rendering
templates = Jinja2Templates(directory="app/templates")

# --- Constants ---
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50 MB per file
ALLOWED_PDF_TYPES = ["application/pdf"]
ALLOWED_DOCX_TYPES = [
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"  # .docx
]


# --- Helpers ---
def validate_file_size(file_bytes: bytes, max_size: int = MAX_FILE_SIZE):
    """Raise if file exceeds max size."""
    if len(file_bytes) > max_size:
        raise HTTPException(
            status_code=413, detail=f"File too large. Max size is {max_size // (1024*1024)} MB."
        )


def validate_pdf(file: UploadFile):
    """Validate PDF MIME type."""
    if file.content_type not in ALLOWED_PDF_TYPES:
        raise HTTPException(
            status_code=400, detail=f"Invalid file type: {file.filename}. Expected PDF."
        )


def validate_docx(file: UploadFile):
    """Validate DOCX MIME type."""
    if file.content_type not in ALLOWED_DOCX_TYPES:
        raise HTTPException(
            status_code=400, detail=f"Invalid file type: {file.filename}. Expected DOCX."
        )


# --- Routes ---
@app.get("/", response_class=HTMLResponse)
async def home(request: Request):
    """Render the main UI."""
    return templates.TemplateResponse("index.html", {"request": request})


@app.post("/api/merge")
async def merge_pdfs(files: List[UploadFile] = File(...)):
    """
    Merge multiple PDFs into one.
    All processing is done in-memory using BytesIO.
    """
    if len(files) < 2:
        raise HTTPException(status_code=400, detail="Provide at least 2 PDF files to merge.")

    # Use first file's name for output
    first_filename = files[0].filename or "document"
    base_name = first_filename.rsplit(".", 1)[0] if "." in first_filename else first_filename
    output_filename = f"{base_name}-merged.pdf"

    writer = PdfWriter()

    for file in files:
        validate_pdf(file)
        content = await file.read()
        validate_file_size(content)
        
        # Read PDF from in-memory bytes
        pdf_stream = io.BytesIO(content)
        reader = PdfReader(pdf_stream)
        for page in reader.pages:
            writer.add_page(page)

    # Write merged PDF to in-memory buffer
    output_buffer = io.BytesIO()
    writer.write(output_buffer)
    output_buffer.seek(0)

    return StreamingResponse(
        output_buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={output_filename}"}
    )


@app.post("/api/split")
async def split_pdf(file: UploadFile = File(...), pages: str = Form(...)):
    """
    Split PDF by page ranges.
    Format: "1-3,5,7-10" extracts pages 1-3, 5, and 7-10.
    Returns a ZIP with individual PDFs for each range.
    """
    validate_pdf(file)
    content = await file.read()
    validate_file_size(content)

    # Use original filename for output
    original_filename = file.filename or "document"
    base_name = original_filename.rsplit(".", 1)[0] if "." in original_filename else original_filename
    output_filename = f"{base_name}-split.zip"

    pdf_stream = io.BytesIO(content)
    reader = PdfReader(pdf_stream)
    total_pages = len(reader.pages)

    # Parse page ranges
    ranges = []
    for part in pages.split(","):
        part = part.strip()
        if "-" in part:
            start, end = part.split("-", 1)
            start = int(start.strip())
            end = int(end.strip())
            if start < 1 or end > total_pages or start > end:
                raise HTTPException(status_code=400, detail=f"Invalid range: {part}")
            ranges.append((start, end))
        else:
            p = int(part)
            if p < 1 or p > total_pages:
                raise HTTPException(status_code=400, detail=f"Invalid page: {p}")
            ranges.append((p, p))

    # Create ZIP with split PDFs in-memory
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for i, (start, end) in enumerate(ranges):
            writer = PdfWriter()
            for p in range(start - 1, end):  # 0-indexed
                writer.add_page(reader.pages[p])
            
            pdf_out = io.BytesIO()
            writer.write(pdf_out)
            pdf_out.seek(0)
            zf.writestr(f"pages_{start}-{end}.pdf", pdf_out.read())

    zip_buffer.seek(0)
    return StreamingResponse(
        zip_buffer,
        media_type="application/zip",
        headers={"Content-Disposition": f"attachment; filename={output_filename}"}
    )


@app.post("/api/compress")
async def compress_pdf(file: UploadFile = File(...)):
    """
    Compress PDF to reduce file size.
    Uses PyMuPDF's deflate and garbage collection.
    """
    validate_pdf(file)
    content = await file.read()
    validate_file_size(content)

    # Use original filename for output
    original_filename = file.filename or "document"
    base_name = original_filename.rsplit(".", 1)[0] if "." in original_filename else original_filename
    output_filename = f"{base_name}-compressed.pdf"

    # Open with PyMuPDF from bytes
    doc = fitz.open(stream=content, filetype="pdf")

    # Write compressed PDF to in-memory buffer
    output_buffer = io.BytesIO()
    doc.save(
        output_buffer,
        garbage=4,  # Maximum garbage collection
        deflate=True,  # Compress streams
        clean=True,  # Clean content streams
    )
    doc.close()
    output_buffer.seek(0)

    return StreamingResponse(
        output_buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={output_filename}"}
    )


@app.post("/api/pdf-to-jpg")
async def pdf_to_jpg(file: UploadFile = File(...)):
    """
    Convert each PDF page to a JPG image.
    Returns a ZIP containing all images.
    """
    validate_pdf(file)
    content = await file.read()
    validate_file_size(content)

    # Use original filename for output
    original_filename = file.filename or "document"
    base_name = original_filename.rsplit(".", 1)[0] if "." in original_filename else original_filename
    output_filename = f"{base_name}-images.zip"

    doc = fitz.open(stream=content, filetype="pdf")

    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for i, page in enumerate(doc):
            # Render page to pixmap (image) at 150 DPI
            pix = page.get_pixmap(dpi=150)
            img_bytes = pix.tobytes("jpeg")
            zf.writestr(f"page_{i + 1}.jpg", img_bytes)

    doc.close()
    zip_buffer.seek(0)

    return StreamingResponse(
        zip_buffer,
        media_type="application/zip",
        headers={"Content-Disposition": f"attachment; filename={output_filename}"}
    )


@app.post("/api/pdf-to-word")
async def pdf_to_word(file: UploadFile = File(...)):
    """
    Convert PDF to DOCX with high fidelity.
    Preserves: text, images, tables, alignment, fonts, colors, bold/italic.
    """
    validate_pdf(file)
    content = await file.read()
    validate_file_size(content)

    # Use original filename for output
    original_filename = file.filename or "document"
    base_name = original_filename.rsplit(".", 1)[0] if "." in original_filename else original_filename
    output_filename = f"{base_name}.docx"

    pdf_doc = fitz.open(stream=content, filetype="pdf")
    word_doc = Document()

    for page_num, page in enumerate(pdf_doc):
        if page_num > 0:
            word_doc.add_page_break()
        
        page_width = page.rect.width
        page_height = page.rect.height
        
        # Dynamic margin tolerance (5% of page width)
        margin_tolerance = page_width * 0.05
        min_margin_for_center = page_width * 0.1  # At least 10% margin to be centered

        # Collect all content elements with Y positions
        elements = []  # (y_pos, type, data)
        
        # --- Extract Tables ---
        tables = page.find_tables()
        table_rects = []
        
        if tables and len(tables.tables) > 0:
            for table in tables:
                table_rects.append(table.bbox)
                table_data = table.extract()
                if table_data and len(table_data) > 0:
                    table_data = [row for row in table_data if any(cell and cell.strip() for cell in row)]
                    if table_data:
                        y_pos = table.bbox[1]
                        elements.append((y_pos, "table", table_data))

        # --- Extract Images ---
        image_list = page.get_images(full=True)
        for img_index, img_info in enumerate(image_list):
            xref = img_info[0]
            try:
                base_image = pdf_doc.extract_image(xref)
                image_bytes = base_image["image"]
                image_ext = base_image["ext"]
                
                # Get image position on page
                img_rects = page.get_image_rects(xref)
                if img_rects:
                    img_rect = img_rects[0]
                    y_pos = img_rect.y0
                    # Store image data with position info
                    elements.append((y_pos, "image", {
                        "bytes": image_bytes,
                        "ext": image_ext,
                        "width": img_rect.width,
                        "height": img_rect.height,
                        "x": img_rect.x0,
                        "page_width": page_width
                    }))
            except Exception:
                pass  # Skip problematic images

        # --- Extract Text Blocks ---
        blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
        
        for block in blocks:
            if block.get("type") != 0:  # Skip image blocks (handled separately)
                continue
            
            block_rect = fitz.Rect(block.get("bbox", (0, 0, 0, 0)))
            
            # Skip if inside table area
            in_table = any(fitz.Rect(tr).intersects(block_rect) for tr in table_rects)
            if in_table:
                continue
            
            # Process each line
            for line in block.get("lines", []):
                line_bbox = line.get("bbox", [0, 0, 0, 0])
                y_pos = line_bbox[1]
                x_start = line_bbox[0]
                x_end = line_bbox[2]
                
                # Collect span info for rich text
                spans_info = []
                for span in line.get("spans", []):
                    text = span.get("text", "")
                    if not text:
                        continue
                    
                    font_size = span.get("size", 12)
                    font_name = span.get("font", "").lower()
                    color_int = span.get("color", 0)
                    
                    # Parse color (BGR integer to RGB)
                    r = (color_int >> 16) & 0xFF
                    g = (color_int >> 8) & 0xFF
                    b = color_int & 0xFF
                    
                    is_bold = "bold" in font_name or "heavy" in font_name or "black" in font_name
                    is_italic = "italic" in font_name or "oblique" in font_name
                    
                    spans_info.append({
                        "text": text,
                        "size": font_size,
                        "bold": is_bold,
                        "italic": is_italic,
                        "color": (r, g, b)
                    })
                
                if not spans_info:
                    continue
                
                # Determine alignment
                left_margin = x_start
                right_margin = page_width - x_end
                
                alignment = "left"
                if abs(left_margin - right_margin) < margin_tolerance and left_margin > min_margin_for_center:
                    alignment = "center"
                elif right_margin < margin_tolerance and left_margin > min_margin_for_center:
                    alignment = "right"
                
                # Get max font size for heading detection
                max_size = max(s["size"] for s in spans_info)
                
                elements.append((y_pos, "text", {
                    "spans": spans_info,
                    "alignment": alignment,
                    "max_size": max_size
                }))
        
        # Sort by Y position
        elements.sort(key=lambda x: x[0])
        
        # --- Output elements in order ---
        for y_pos, elem_type, data in elements:
            if elem_type == "table":
                table_data = data
                num_cols = max(len(row) for row in table_data) if table_data else 1
                word_table = word_doc.add_table(rows=len(table_data), cols=num_cols)
                word_table.style = 'Table Grid'
                
                for i, row in enumerate(table_data):
                    for j, cell in enumerate(row):
                        if j < num_cols and cell:
                            word_table.rows[i].cells[j].text = cell.strip()
                
                word_doc.add_paragraph()
                
            elif elem_type == "image":
                try:
                    img_stream = io.BytesIO(data["bytes"])
                    # Scale image to fit page width (max 6 inches)
                    img_width_inches = min(data["width"] / 72, 6)
                    
                    para = word_doc.add_paragraph()
                    run = para.add_run()
                    run.add_picture(img_stream, width=Inches(img_width_inches))
                    
                    # Center image if it was centered in PDF
                    img_x = data["x"]
                    img_right = data["page_width"] - (img_x + data["width"])
                    if abs(img_x - img_right) < margin_tolerance:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    pass  # Skip problematic images
                
            elif elem_type == "text":
                spans = data["spans"]
                alignment = data["alignment"]
                max_size = data["max_size"]
                
                # Create paragraph
                para = word_doc.add_paragraph()
                
                # Apply alignment
                if alignment == "center":
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif alignment == "right":
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                # Add runs with formatting
                for span_info in spans:
                    run = para.add_run(span_info["text"])
                    run.bold = span_info["bold"]
                    run.italic = span_info["italic"]
                    run.font.size = Pt(span_info["size"])
                    
                    # Apply color if not black
                    r, g, b = span_info["color"]
                    if not (r == 0 and g == 0 and b == 0):
                        run.font.color.rgb = RGBColor(r, g, b)

    pdf_doc.close()

    output_buffer = io.BytesIO()
    word_doc.save(output_buffer)
    output_buffer.seek(0)

    return StreamingResponse(
        output_buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={output_filename}"}
    )


@app.post("/api/word-to-pdf")
async def word_to_pdf(file: UploadFile = File(...)):
    """
    Convert DOCX to PDF with high fidelity.
    Preserves: alignment, images, tables with borders, bullets, bold/italic, colors.
    """
    validate_docx(file)
    content = await file.read()
    validate_file_size(content)

    original_filename = file.filename or "document"
    base_name = original_filename.rsplit(".", 1)[0] if "." in original_filename else original_filename
    output_filename = f"{base_name}.pdf"

    docx_stream = io.BytesIO(content)
    word_doc = Document(docx_stream)

    output_buffer = io.BytesIO()
    c = canvas.Canvas(output_buffer, pagesize=letter)
    width, height = letter
    
    left_margin = 50
    right_margin = width - 50
    content_width = right_margin - left_margin

    y = height - 50
    line_height = 16
    
    def check_page(needed_space=50):
        nonlocal y
        if y < needed_space:
            c.showPage()
            y = height - 50
    
    def get_alignment_x(text, font_name, font_size, alignment):
        """Calculate X position based on alignment."""
        text_width = c.stringWidth(text, font_name, font_size)
        if alignment == "center":
            return left_margin + (content_width - text_width) / 2
        elif alignment == "right":
            return right_margin - text_width
        return left_margin  # left alignment
    
    def draw_text_aligned(text, font_name="Helvetica", font_size=11, bold=False, italic=False, 
                          alignment="left", color=None, indent=0):
        nonlocal y
        if not text.strip():
            y -= line_height // 2
            return
        
        # Build font name
        if bold and italic:
            actual_font = font_name + "-BoldOblique"
        elif bold:
            actual_font = font_name + "-Bold"
        elif italic:
            actual_font = font_name + "-Oblique"
        else:
            actual_font = font_name
        
        try:
            c.setFont(actual_font, font_size)
        except:
            try:
                c.setFont(font_name + "-Bold" if bold else font_name, font_size)
            except:
                c.setFont("Helvetica", font_size)
        
        # Set color
        if color:
            c.setFillColorRGB(color[0]/255, color[1]/255, color[2]/255)
        else:
            c.setFillColorRGB(0, 0, 0)
        
        effective_left = left_margin + indent
        effective_width = content_width - indent
        
        # Word wrapping
        words = text.split()
        line = ""
        for word in words:
            test_line = f"{line} {word}".strip()
            try:
                line_width = c.stringWidth(test_line, actual_font, font_size)
            except:
                line_width = len(test_line) * font_size * 0.5
            
            if line_width < effective_width:
                line = test_line
            else:
                if line:
                    # Draw with alignment
                    if alignment == "center":
                        x = effective_left + (effective_width - c.stringWidth(line, actual_font, font_size)) / 2
                    elif alignment == "right":
                        x = right_margin - c.stringWidth(line, actual_font, font_size)
                    else:
                        x = effective_left
                    
                    c.drawString(x, y, line)
                    y -= line_height
                    check_page()
                line = word
        
        # Draw remaining text
        if line:
            try:
                line_str_width = c.stringWidth(line, actual_font, font_size)
            except:
                line_str_width = len(line) * font_size * 0.5
                
            if alignment == "center":
                x = effective_left + (effective_width - line_str_width) / 2
            elif alignment == "right":
                x = right_margin - line_str_width
            else:
                x = effective_left
            
            c.drawString(x, y, line)
            y -= line_height
        
        check_page()
        c.setFillColorRGB(0, 0, 0)  # Reset to black

    def draw_table_with_borders(table):
        nonlocal y
        
        # Calculate table dimensions
        rows_data = []
        max_cols = 0
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells]
            rows_data.append(row_cells)
            max_cols = max(max_cols, len(row_cells))
        
        if not rows_data:
            return
        
        # Check if enough space for at least header
        check_page(100)
        
        # Calculate column widths
        col_width = content_width / max_cols
        row_height = 20
        table_start_y = y
        
        c.setFont("Helvetica", 9)
        
        for row_idx, row_cells in enumerate(rows_data):
            check_page(row_height + 20)
            
            # Draw row
            x = left_margin
            for col_idx, cell_text in enumerate(row_cells):
                # Draw cell border
                c.rect(x, y - row_height + 4, col_width, row_height)
                
                # Draw cell text (truncate if too long)
                max_chars = int(col_width / 5)
                display_text = cell_text[:max_chars] if len(cell_text) > max_chars else cell_text
                c.drawString(x + 3, y - 2, display_text)
                
                x += col_width
            
            y -= row_height
        
        y -= 10  # Space after table

    def draw_image(image_bytes, alignment="left"):
        nonlocal y
        try:
            img = Image.open(io.BytesIO(image_bytes))
            img_width, img_height = img.size
            
            # Scale to fit content width (max 5 inches)
            max_width = min(content_width, 5 * 72)
            scale = min(max_width / img_width, 1)
            
            draw_width = img_width * scale
            draw_height = img_height * scale
            
            # Check page space
            check_page(draw_height + 20)
            
            # Calculate X based on alignment
            if alignment == "center":
                x = left_margin + (content_width - draw_width) / 2
            elif alignment == "right":
                x = right_margin - draw_width
            else:
                x = left_margin
            
            # Save image to temp buffer for ReportLab
            img_buffer = io.BytesIO()
            img.save(img_buffer, format='PNG')
            img_buffer.seek(0)
            
            from reportlab.lib.utils import ImageReader
            img_reader = ImageReader(img_buffer)
            
            c.drawImage(img_reader, x, y - draw_height, draw_width, draw_height)
            y -= draw_height + 10
            
        except Exception:
            pass  # Skip problematic images

    def get_para_alignment(para):
        """Get paragraph alignment from Word document."""
        try:
            alignment = para.alignment
            if alignment is not None:
                if alignment == WD_ALIGN_PARAGRAPH.CENTER:
                    return "center"
                elif alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                    return "right"
                elif alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                    return "left"  # Treat as left for simplicity
        except (ValueError, KeyError):
            # Handle unknown alignment values like 'start'
            pass
        return "left"
    
    def get_run_color(run):
        """Extract color from run if set."""
        try:
            if run.font.color and run.font.color.rgb:
                rgb = run.font.color.rgb
                return (rgb[0], rgb[1], rgb[2])
        except:
            pass
        return None

    # Process document body elements in order
    for element in word_doc.element.body:
        if element.tag.endswith('p'):  # Paragraph
            for para in word_doc.paragraphs:
                if para._element == element:
                    text = para.text
                    style_name = para.style.name if para.style else ""
                    alignment = get_para_alignment(para)
                    
                    # Check for bullet/list
                    is_list = False
                    indent = 0
                    prefix = ""
                    
                    # Check if it's a list item by looking at numPr element
                    numPr = para._element.find('.//' + qn('w:numPr'))
                    if numPr is not None:
                        is_list = True
                        indent = 20
                        prefix = "• "
                    
                    # Handle images in paragraph
                    for run in para.runs:
                        # Check for inline images
                        drawing_elements = run._element.findall('.//' + qn('a:blip'))
                        for blip in drawing_elements:
                            rId = blip.get(qn('r:embed'))
                            if rId:
                                try:
                                    image_part = word_doc.part.related_parts[rId]
                                    draw_image(image_part.blob, alignment)
                                except:
                                    pass
                    
                    if not text.strip():
                        y -= line_height // 2
                        continue
                    
                    # Determine formatting
                    if "Heading" in style_name:
                        level = 1
                        try:
                            level = int(style_name.split()[-1])
                        except:
                            pass
                        font_size = 18 - (level * 2) if level <= 3 else 11
                        draw_text_aligned(text, font_size=max(font_size, 12), bold=True, 
                                         alignment=alignment)
                        y -= 4
                    elif "Title" in style_name:
                        draw_text_aligned(text, font_size=20, bold=True, alignment=alignment)
                        y -= 6
                    else:
                        # Process runs for mixed formatting
                        has_runs = len(para.runs) > 0
                        
                        if has_runs and len(para.runs) > 1:
                            # Mixed formatting - process run by run on same line
                            # Simplified: just use paragraph-level detection
                            is_bold = any(run.bold for run in para.runs if run.bold)
                            is_italic = any(run.italic for run in para.runs if run.italic)
                            color = None
                            for run in para.runs:
                                run_color = get_run_color(run)
                                if run_color:
                                    color = run_color
                                    break
                            
                            display_text = prefix + text if is_list else text
                            draw_text_aligned(display_text, bold=is_bold, italic=is_italic, 
                                            alignment=alignment, color=color, indent=indent)
                        else:
                            is_bold = para.runs[0].bold if para.runs else False
                            is_italic = para.runs[0].italic if para.runs else False
                            color = get_run_color(para.runs[0]) if para.runs else None
                            
                            display_text = prefix + text if is_list else text
                            draw_text_aligned(display_text, bold=is_bold, italic=is_italic, 
                                            alignment=alignment, color=color, indent=indent)
                    break
        
        elif element.tag.endswith('tbl'):  # Table
            for table in word_doc.tables:
                if table._element == element:
                    draw_table_with_borders(table)
                    break
    
    # Fallback if no content
    if y == height - 50:
        for para in word_doc.paragraphs:
            text = para.text
            if text.strip():
                alignment = get_para_alignment(para)
                is_bold = any(run.bold for run in para.runs if run.bold)
                draw_text_aligned(text, bold=is_bold, alignment=alignment)
        
        for table in word_doc.tables:
            draw_table_with_borders(table)

    c.save()
    output_buffer.seek(0)

    return StreamingResponse(
        output_buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={output_filename}"}
    )


# --- Health Check ---
@app.get("/health")
async def health():
    return {"status": "ok"}
