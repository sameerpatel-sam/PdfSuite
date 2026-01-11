import io
import zipfile
from typing import List
from pypdf import PdfReader, PdfWriter
from pdfminer.high_level import extract_text
from docx import Document
from PIL import Image
from pdf2image import convert_from_bytes
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas


def merge_pdfs(buffers: List[io.BytesIO]) -> bytes:
    """Merge multiple PDFs into one using PdfWriter."""
    writer = PdfWriter()
    for b in buffers:
        b.seek(0)
        reader = PdfReader(b)
        for page in reader.pages:
            writer.add_page(page)
    out = io.BytesIO()
    writer.write(out)
    writer.close()
    return out.getvalue()


def _parse_ranges(ranges: str, total_pages: int) -> List[List[int]]:
    # returns list of page indices (0-based) groups
    groups: List[List[int]] = []
    parts = [p.strip() for p in ranges.split(',') if p.strip()]
    for part in parts:
        if '-' in part:
            a, b = part.split('-', 1)
            start = int(a) if a else 1
            end = int(b) if b else total_pages
            if start < 1 or end < start:
                raise ValueError("Invalid range: " + part)
            groups.append([i-1 for i in range(start, min(end, total_pages) + 1)])
        else:
            idx = int(part)
            if idx < 1 or idx > total_pages:
                raise ValueError("Invalid page: " + part)
            groups.append([idx-1])
    if not groups:
        raise ValueError("No valid ranges provided")
    return groups


def split_pdf_by_ranges(buffer: io.BytesIO, ranges: str) -> bytes:
    buffer.seek(0)
    reader = PdfReader(buffer)
    total = len(reader.pages)
    groups = _parse_ranges(ranges, total)

    out_zip = io.BytesIO()
    with zipfile.ZipFile(out_zip, mode='w', compression=zipfile.ZIP_DEFLATED) as zf:
        for idx, group in enumerate(groups, start=1):
            writer = PdfWriter()
            for p in group:
                writer.add_page(reader.pages[p])
            pdf_bytes = io.BytesIO()
            writer.write(pdf_bytes)
            writer.close()
            pdf_bytes.seek(0)
            zf.writestr(f"part-{idx}.pdf", pdf_bytes.read())
    return out_zip.getvalue()


def compress_pdf_basic(buffer: io.BytesIO) -> bytes:
    buffer.seek(0)
    reader = PdfReader(buffer)
    writer = PdfWriter()
    for page in reader.pages:
        # Basic content stream compression; does not re-encode images
        try:
            page.compress_content_streams()
        except Exception:
            pass
        writer.add_page(page)
    out = io.BytesIO()
    writer.write(out)
    writer.close()
    return out.getvalue()


def pdf_to_jpg_zip(data: bytes, dpi: int = 150, quality: int = 85) -> bytes:
    images = convert_from_bytes(data, dpi=dpi)
    out_zip = io.BytesIO()
    with zipfile.ZipFile(out_zip, mode='w', compression=zipfile.ZIP_DEFLATED) as zf:
        for i, img in enumerate(images, start=1):
            b = io.BytesIO()
            rgb = img.convert('RGB')
            rgb.save(b, format='JPEG', quality=quality, optimize=True)
            b.seek(0)
            zf.writestr(f"page-{i}.jpg", b.read())
    return out_zip.getvalue()


def pdf_to_docx_bytes(data: bytes) -> bytes:
    # Best-effort: extract text and write to DOCX; layout not preserved.
    text = extract_text(io.BytesIO(data)) or ""
    doc = Document()
    doc.add_heading('Converted from PDF', level=1)
    for line in text.splitlines():
        doc.add_paragraph(line)
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def docx_to_pdf_bytes(data: bytes) -> bytes:
    """
    Convert DOCX to PDF using reportlab.
    Extracts text from DOCX and renders to PDF. Complex formatting not preserved.
    """
    docx_stream = io.BytesIO(data)
    word_doc = Document(docx_stream)
    paragraphs = [p.text for p in word_doc.paragraphs]

    out = io.BytesIO()
    c = canvas.Canvas(out, pagesize=letter)
    width, height = letter
    y = height - 50
    line_height = 14

    for para in paragraphs:
        if not para.strip():
            y -= line_height
            continue
        words = para.split()
        line = ""
        for word in words:
            test_line = f"{line} {word}".strip()
            if c.stringWidth(test_line, "Helvetica", 11) < width - 100:
                line = test_line
            else:
                c.drawString(50, y, line)
                y -= line_height
                line = word
                if y < 50:
                    c.showPage()
                    y = height - 50
        if line:
            c.drawString(50, y, line)
            y -= line_height
        if y < 50:
            c.showPage()
            y = height - 50

    c.save()
    return out.getvalue()
